from flask import Flask, request, render_template, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import json
from io import BytesIO

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'doc'}

# Tạo thư mục uploads nếu chưa có
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def remove_bullets(paragraph):
    """Bỏ các bullets thừa ở đầu dòng"""
    text = paragraph.text.strip()
    original_text = text
    # Bỏ các ký tự bullets phổ biến
    bullets = ['•', '·', '○', '▪', '▫', '-', '–', '—']
    for bullet in bullets:
        if text.startswith(bullet):
            text = text[1:].strip()
    # Bỏ số thứ tự đầu dòng (1., 2., a., b., etc.)
    text = re.sub(r'^\d+[\.\)]\s*', '', text)
    text = re.sub(r'^[a-z][\.\)]\s*', '', text)
    text = re.sub(r'^[ivx]+[\.\)]\s*', '', text, flags=re.IGNORECASE)
    
    # Chỉ cập nhật nếu có thay đổi
    if text != original_text:
        # Xóa toàn bộ runs và thêm lại với nội dung mới
        paragraph.clear()
        if text:  # Chỉ thêm nếu còn text
            paragraph.add_run(text)
    
    return text

def standardize_font(paragraph, is_title=False):
    """Chuẩn hóa font theo nghị định 30/2020/NĐ-CP"""
    # Font chữ: Times New Roman
    # Cỡ chữ: 13-14 cho nội dung, 14-16 cho tiêu đề
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if is_title:
            run.font.size = Pt(14)
            run.font.bold = True
        else:
            run.font.size = Pt(13)
            run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)  # Đen

def standardize_paragraph_spacing(paragraph, is_title=False):
    """Chuẩn hóa khoảng cách đoạn văn"""
    # Khoảng cách dòng: 1.2-1.3 lines
    paragraph.paragraph_format.line_spacing = 1.25
    # Khoảng cách trước/sau: tối thiểu 6pt
    if is_title:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

def standardize_alignment(paragraph, is_title=False):
    """Chuẩn hóa căn lề"""
    if is_title:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def is_title_paragraph(paragraph, index, total):
    """Xác định xem đoạn văn có phải là tiêu đề không"""
    text = paragraph.text.strip()
    # Nếu là đoạn đầu tiên và ngắn (< 100 ký tự), có thể là tiêu đề
    if index == 0 and len(text) < 100:
        return True
    # Nếu có font lớn hoặc bold
    if paragraph.runs:
        run = paragraph.runs[0]
        if run.font.size and run.font.size.pt >= 14:
            return True
        if run.bold:
            return True
    # Nếu là dòng ngắn và có chữ in hoa
    if len(text) < 80 and text.isupper():
        return True
    return False

def standardize_document(doc):
    """Chuẩn hóa toàn bộ document theo nghị định 30/2020/NĐ-CP"""
    # Chuẩn hóa margins: trên 2cm, dưới 2cm, trái 3cm, phải 2cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Chuẩn hóa từng paragraph
    paragraphs = doc.paragraphs
    total = len(paragraphs)
    
    for idx, para in enumerate(paragraphs):
        # Bỏ bullets thừa (hàm này đã tự xử lý clear và add_run)
        text = remove_bullets(para)
        
        # Xác định xem có phải tiêu đề không
        is_title = is_title_paragraph(para, idx, total)
        
        # Chuẩn hóa font, spacing, alignment
        standardize_font(para, is_title)
        standardize_paragraph_spacing(para, is_title)
        standardize_alignment(para, is_title)
        
        # Bỏ bullets formatting
        para.paragraph_format.first_line_indent = None
        if para.paragraph_format.left_indent:
            para.paragraph_format.left_indent = None
    
    # Xử lý tables nếu có
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    remove_bullets(para)  # Hàm này đã tự xử lý clear và add_run
                    standardize_font(para)
                    standardize_paragraph_spacing(para)
                    standardize_alignment(para)

def apply_custom_adjustments(doc, adjustments):
    """Áp dụng các điều chỉnh từ form"""
    if not adjustments:
        return
    
    data = json.loads(adjustments) if isinstance(adjustments, str) else adjustments
    
    # Điều chỉnh font size
    if 'fontSize' in data:
        size = int(data['fontSize'])
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.size = Pt(size)
    
    # Điều chỉnh line spacing
    if 'lineSpacing' in data:
        spacing = float(data['lineSpacing'])
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = spacing
    
    # Điều chỉnh margins
    if 'margins' in data:
        margins = data['margins']
        for section in doc.sections:
            if 'top' in margins:
                section.top_margin = Cm(float(margins['top']))
            if 'bottom' in margins:
                section.bottom_margin = Cm(float(margins['bottom']))
            if 'left' in margins:
                section.left_margin = Cm(float(margins['left']))
            if 'right' in margins:
                section.right_margin = Cm(float(margins['right']))

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            # Đọc và chuẩn hóa document
            doc = docx.Document(filepath)
            
            # Lưu bản gốc để so sánh
            original_text = []
            for para in doc.paragraphs:
                original_text.append(para.text)
            
            # Chuẩn hóa document
            standardize_document(doc)
            
            # Lấy nội dung sau chuẩn hóa
            standardized_text = []
            total_paras = len(doc.paragraphs)
            for idx, para in enumerate(doc.paragraphs):
                standardized_text.append({
                    'text': para.text,
                    'is_title': is_title_paragraph(para, idx, total_paras)
                })
            
            # Lưu file đã chuẩn hóa tạm thời
            temp_filename = f"standardized_{filename}"
            temp_filepath = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
            doc.save(temp_filepath)
            
            return jsonify({
                'success': True,
                'filename': filename,
                'temp_filename': temp_filename,
                'content': standardized_text,
                'message': 'File đã được chuẩn hóa thành công!'
            })
        
        except Exception as e:
            return jsonify({'error': f'Lỗi xử lý file: {str(e)}'}), 500
    
    return jsonify({'error': 'File không hợp lệ. Vui lòng upload file .docx hoặc .doc'}), 400

@app.route('/adjust', methods=['POST'])
def adjust_document():
    try:
        data = request.json
        filename = data.get('filename')
        adjustments = data.get('adjustments', {})
        
        if not filename:
            return jsonify({'error': 'Missing filename'}), 400
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"standardized_{filename}")
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        # Đọc document
        doc = docx.Document(filepath)
        
        # Áp dụng điều chỉnh
        apply_custom_adjustments(doc, adjustments)
        
        # Lưu lại
        doc.save(filepath)
        
        # Lấy nội dung mới
        content = []
        total_paras = len(doc.paragraphs)
        for idx, para in enumerate(doc.paragraphs):
            content.append({
                'text': para.text,
                'is_title': is_title_paragraph(para, idx, total_paras)
            })
        
        return jsonify({
            'success': True,
            'content': content,
            'message': 'Đã áp dụng điều chỉnh thành công!'
        })
    
    except Exception as e:
        return jsonify({'error': f'Lỗi điều chỉnh: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"standardized_{filename}")
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f"Chuan_hoa_{filename}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': f'Lỗi download: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)

